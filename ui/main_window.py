"""主窗口 — 浅色简洁风，支持本地文件和视频链接"""

import os
import json
import threading
import base64
from io import BytesIO
from PySide6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QFrame,
    QLabel, QPushButton, QTabWidget, QTextEdit, QProgressBar,
    QFileDialog, QApplication, QLineEdit, QStackedWidget, QDialog,
    QScrollArea
)
from PySide6.QtCore import Qt, Signal, Slot, QMetaObject, Q_ARG, QSize
from PySide6.QtGui import QDragEnterEvent, QDropEvent, QImage, QPixmap


class ClickableImageLabel(QLabel):
    """可点击放大的图片标签"""
    clicked = Signal(str)

    def __init__(self, image_path, parent=None):
        super().__init__(parent)
        self.image_path = image_path
        self.setCursor(Qt.PointingHandCursor)
        self.setStyleSheet("QLabel:hover { border: 2px solid #007AFF; border-radius: 6px; }")

    def mousePressEvent(self, event):
        self.clicked.emit(self.image_path)


class ImageDialog(QDialog):
    """图片放大预览对话框"""

    def __init__(self, image_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle("场景截图")
        self.setMinimumSize(640, 480)
        self.resize(800, 600)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        img_label = QLabel()
        img_label.setAlignment(Qt.AlignCenter)
        pixmap = QPixmap(image_path)
        if not pixmap.isNull():
            # 自适应窗口大小，但不超过原图
            scaled = pixmap.scaled(780, 580, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            img_label.setPixmap(scaled)
        scroll.setWidget(img_label)
        layout.addWidget(scroll)

        btn = QPushButton("关闭")
        btn.clicked.connect(self.accept)
        btn.setStyleSheet(
            "QPushButton { background-color: #007AFF; color: white; border: none; "
            "border-radius: 8px; padding: 8px 20px; font-weight: 500; }"
        )
        layout.addWidget(btn, alignment=Qt.AlignCenter)


class ContentExpandDialog(QDialog):
    """内容放大查看对话框"""

    def __init__(self, title, content, is_markdown=True, parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setMinimumSize(800, 600)
        self.resize(1000, 750)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        text_widget = QTextEdit()
        text_widget.setReadOnly(True)
        if is_markdown:
            text_widget.setMarkdown(content)
        else:
            text_widget.setHtml(content)
        text_widget.setStyleSheet(
            "QTextEdit { background-color: #FFFFFF; border: none; "
            "font-size: 15px; line-height: 1.6; padding: 12px; }"
        )
        scroll.setWidget(text_widget)
        layout.addWidget(scroll)

        btn_row = QHBoxLayout()
        btn_row.addStretch()
        btn = QPushButton("关闭")
        btn.clicked.connect(self.accept)
        btn.setStyleSheet(
            "QPushButton { background-color: #007AFF; color: white; border: none; "
            "border-radius: 8px; padding: 8px 24px; font-weight: 500; font-size: 14px; }"
        )
        btn_row.addWidget(btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

from core.pipeline import VideoToScriptPipeline, AnalysisResult
from core.downloader import download_video, get_video_info_only, detect_platform
from ui.settings_dialog import SettingsDialog, load_settings, save_settings


class DropZone(QFrame):
    file_dropped = Signal(str)

    def __init__(self):
        super().__init__()
        self.setAcceptDrops(True)
        self.setObjectName("dropZone")
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignCenter)
        layout.setSpacing(8)
        icon = QLabel("🎬")
        icon.setObjectName("dropIcon")
        icon.setAlignment(Qt.AlignCenter)
        layout.addWidget(icon)
        text = QLabel("拖拽视频文件到这里")
        text.setObjectName("dropText")
        text.setAlignment(Qt.AlignCenter)
        layout.addWidget(text)
        hint = QLabel("或者点击下方按钮选择文件 · 支持 MP4 / MOV / MKV / AVI / WEBM")
        hint.setObjectName("dropHint")
        hint.setAlignment(Qt.AlignCenter)
        hint.setWordWrap(True)
        layout.addWidget(hint)

    def _resolve_path(self, mime):
        """从 MIME 数据解析本地文件路径"""
        if mime.hasUrls():
            url = mime.urls()[0]
            path = url.toLocalFile()
            if not path:
                url_str = url.toString()
                if url_str.startswith("file://"):
                    from urllib.parse import unquote, urlparse
                    path = unquote(urlparse(url_str).path)
            return path
        if mime.hasText():
            return mime.text().strip()
        return ""

    def _is_video(self, path):
        if not path:
            return False
        ext = os.path.splitext(path)[1].lower()
        return ext in (".mp4", ".mov", ".mkv", ".avi", ".webm", ".flv", ".wmv", ".ts")

    def dragEnterEvent(self, event):
        path = self._resolve_path(event.mimeData())
        if self._is_video(path):
            self.setProperty("dragOver", True)
            self.style().unpolish(self)
            self.style().polish(self)
            event.acceptProposedAction()

    def dragMoveEvent(self, event):
        event.acceptProposedAction()

    def dragLeaveEvent(self, event):
        self.setProperty("dragOver", False)
        self.style().unpolish(self)
        self.style().polish(self)

    def dropEvent(self, event):
        self.setProperty("dragOver", False)
        self.style().unpolish(self)
        self.style().polish(self)
        path = self._resolve_path(event.mimeData())
        if path:
            self.file_dropped.emit(path)


class URLInputZone(QFrame):
    """链接输入区域"""
    download_requested = Signal(str)
    preview_requested = Signal(str)

    def __init__(self):
        super().__init__()
        self.setObjectName("urlZone")
        layout = QVBoxLayout(self)
        layout.setSpacing(12)

        header = QHBoxLayout()
        icon = QLabel("🔗")
        icon.setStyleSheet("font-size: 24px;")
        header.addWidget(icon)
        title = QLabel("粘贴视频链接")
        title.setStyleSheet("font-size: 15px; font-weight: 500; color: #1D1D1F;")
        header.addWidget(title)
        header.addStretch()
        self.platform_tag = QLabel("")
        self.platform_tag.setObjectName("platformTag")
        self.platform_tag.setStyleSheet(
            "font-size: 11px; padding: 2px 8px; border-radius: 4px; "
            "background-color: #F0F5FF; color: #007AFF; font-weight: 500;"
        )
        self.platform_tag.setVisible(False)
        header.addWidget(self.platform_tag)
        layout.addLayout(header)

        url_layout = QHBoxLayout()
        url_layout.setSpacing(8)
        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("粘贴抖音/YouTube/B站/快手等视频链接...")
        self.url_input.setObjectName("urlInput")
        self.url_input.setStyleSheet(
            "QLineEdit { border: 1px solid #E5E5EA; border-radius: 8px; "
            "padding: 10px 14px; font-size: 13px; background-color: #FFFFFF; }"
            "QLineEdit:focus { border-color: #007AFF; }"
        )
        self.url_input.textChanged.connect(self._on_url_changed)
        url_layout.addWidget(self.url_input, 1)

        self.preview_btn = QPushButton("预览")
        self.preview_btn.setStyleSheet(
            "QPushButton { background-color: transparent; border: 1px solid #E5E5EA; "
            "border-radius: 8px; padding: 8px 16px; color: #636366; font-weight: 500; }"
            "QPushButton:hover { background-color: #F5F5F7; }"
            "QPushButton:disabled { color: #C7C7CC; }"
        )
        self.preview_btn.setEnabled(False)
        url_layout.addWidget(self.preview_btn)

        self.download_btn = QPushButton("⬇ 下载")
        self.download_btn.setStyleSheet(
            "QPushButton { background-color: #007AFF; color: white; border: none; "
            "border-radius: 8px; padding: 8px 20px; font-weight: 500; }"
            "QPushButton:hover { background-color: #0066D6; }"
            "QPushButton:disabled { background-color: #B0D4FF; color: #E5E5EA; }"
        )
        self.download_btn.setEnabled(False)
        url_layout.addWidget(self.download_btn)
        layout.addLayout(url_layout)

        self.preview_frame = QFrame()
        self.preview_frame.setStyleSheet(
            "QFrame { background-color: #F5F5F7; border-radius: 8px; padding: 8px 12px; }"
        )
        self.preview_frame.setVisible(False)
        pv_layout = QVBoxLayout(self.preview_frame)
        pv_layout.setSpacing(4)
        self.preview_title = QLabel("")
        self.preview_title.setStyleSheet("font-size: 13px; font-weight: 500;")
        pv_layout.addWidget(self.preview_title)
        self.preview_detail = QLabel("")
        self.preview_detail.setStyleSheet("font-size: 11px; color: #86868B;")
        self.preview_detail.setWordWrap(True)
        pv_layout.addWidget(self.preview_detail)
        layout.addWidget(self.preview_frame)

        hint = QLabel("支持 抖音 · YouTube · B站 · 快手 · 西瓜 · 微博 · TikTok · Instagram 等 1000+ 平台")
        hint.setStyleSheet("font-size: 10px; color: #AEAEB2;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

    def _on_url_changed(self, text):
        has_url = bool(text.strip()) and text.strip().startswith("http")
        self.preview_btn.setEnabled(has_url)
        self.download_btn.setEnabled(has_url)
        if has_url:
            platform = detect_platform(text.strip())
            self.platform_tag.setText(platform)
            self.platform_tag.setVisible(True)
        else:
            self.platform_tag.setVisible(False)

    def show_preview(self, title, platform, duration, uploader):
        self.preview_frame.setVisible(True)
        self.preview_title.setText(f"📹 {title}")
        dur_str = f"{int(duration // 60)}:{int(duration % 60):02d}" if duration else "未知"
        self.preview_detail.setText(f"平台: {platform} · 时长: {dur_str} · 作者: {uploader}")


class InputSwitch(QFrame):
    mode_changed = Signal(str)

    def __init__(self):
        super().__init__()
        self.setObjectName("inputSwitch")
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self.local_btn = QPushButton("📁 本地文件")
        self.local_btn.setCheckable(True)
        self.local_btn.setChecked(True)
        layout.addWidget(self.local_btn)

        self.url_btn = QPushButton("🔗 视频链接")
        self.url_btn.setCheckable(True)
        layout.addWidget(self.url_btn)

        self.local_btn.clicked.connect(lambda: self._switch("local"))
        self.url_btn.clicked.connect(lambda: self._switch("url"))

    def _switch(self, mode):
        self.local_btn.setChecked(mode == "local")
        self.url_btn.setChecked(mode == "url")
        active_style = (
            "QPushButton { background-color: #007AFF; color: white; border: none; "
            "border-radius: 6px; padding: 6px 16px; font-weight: 500; font-size: 12px; }"
        )
        inactive_style = (
            "QPushButton { background-color: transparent; color: #86868B; border: none; "
            "border-radius: 6px; padding: 6px 16px; font-weight: 500; font-size: 12px; }"
            "QPushButton:hover { background-color: #F5F5F7; }"
        )
        self.local_btn.setStyleSheet(active_style if mode == "local" else inactive_style)
        self.url_btn.setStyleSheet(active_style if mode == "url" else inactive_style)
        self.mode_changed.emit(mode)


class MainWindow(QMainWindow):
    # 线程安全的信号，用于从子线程更新 UI
    _progress_signal = Signal(int, str)       # (percent, step_text)
    _analyze_done_signal = Signal()
    _download_progress_signal = Signal(str)   # msg
    _preview_done_signal = Signal(object)     # result
    _download_done_signal = Signal(object)    # result
    _batch_url_progress_signal = Signal(int, int, str)  # (current, total, msg) 多集批量进度
    _batch_url_done_signal = Signal()                   # 多集批量完成

    def __init__(self):
        super().__init__()
        self.setWindowTitle("Video to Script")
        self.setMinimumSize(860, 640)
        self.resize(960, 700)
        self.setAcceptDrops(True)
        self.settings = load_settings()
        self.result = None
        self.batch_results = []          # 批量分析结果列表（持久化，不清空）
        self.current_episode_idx = 0     # 当前查看的集数索引
        self.view_mode = "single"        # "single" | "batch" | "summary"
        self.cached_summary_md = ""      # 全局总结缓存（生成后保存在界面中）
        self.is_analyzing = False
        self.is_downloading = False
        self.is_batch_url_analyzing = False  # 多集链接批量分析进行中标记
        self.video_path = ""
        self.video_source = ""
        self._setup_ui()
        self._connect_signals()

    def _setup_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # 顶栏
        top_bar = QFrame()
        top_bar.setObjectName("topBar")
        top_bar.setFixedHeight(56)
        top_layout = QHBoxLayout(top_bar)
        top_layout.setContentsMargins(20, 0, 20, 0)
        title = QLabel("Video to Script")
        title.setObjectName("appTitle")
        top_layout.addWidget(title)
        subtitle = QLabel("视频转剧本")
        subtitle.setObjectName("appSubtitle")
        top_layout.addWidget(subtitle)
        top_layout.addStretch()
        self.settings_btn = QPushButton("⚙ 设置")
        self.settings_btn.setObjectName("settingsBtn")
        top_layout.addWidget(self.settings_btn)
        main_layout.addWidget(top_bar)

        # 内容
        content = QWidget()
        content_layout = QVBoxLayout(content)
        content_layout.setContentsMargins(24, 20, 24, 24)
        content_layout.setSpacing(16)

        # 输入模式切换
        self.input_switch = InputSwitch()
        content_layout.addWidget(self.input_switch)

        # 输入区域
        self.input_stack = QStackedWidget()

        # Page 0: 本地文件
        local_page = QWidget()
        local_layout = QHBoxLayout(local_page)
        local_layout.setSpacing(16)
        self.drop_zone = DropZone()
        self.drop_zone.setMinimumHeight(140)
        local_layout.addWidget(self.drop_zone, 3)

        right = QVBoxLayout()
        right.setSpacing(12)
        self.file_info_frame = QFrame()
        self.file_info_frame.setObjectName("fileInfo")
        fi_layout = QVBoxLayout(self.file_info_frame)
        fi_layout.setSpacing(6)
        self.file_name_label = QLabel("尚未选择文件")
        self.file_name_label.setStyleSheet("font-weight: 500; font-size: 14px;")
        fi_layout.addWidget(self.file_name_label)
        self.file_detail_label = QLabel("")
        self.file_detail_label.setStyleSheet("color: #86868B; font-size: 12px;")
        fi_layout.addWidget(self.file_detail_label)
        self.api_status_label = QLabel("")
        self.api_status_label.setStyleSheet("font-size: 11px;")
        fi_layout.addWidget(self.api_status_label)
        fi_layout.addStretch()
        right.addWidget(self.file_info_frame)
        self.browse_btn = QPushButton("📁 选择文件")
        self.browse_btn.setObjectName("settingsBtn")
        right.addWidget(self.browse_btn)
        right.addStretch()
        local_layout.addLayout(right, 2)
        self.input_stack.addWidget(local_page)

        # Page 1: 视频链接
        url_page = QWidget()
        url_layout = QVBoxLayout(url_page)
        url_layout.setSpacing(12)
        self.url_zone = URLInputZone()
        url_layout.addWidget(self.url_zone)

        self.url_file_info = QFrame()
        self.url_file_info.setObjectName("fileInfo")
        ufi_layout = QVBoxLayout(self.url_file_info)
        ufi_layout.setSpacing(6)
        self.url_file_name = QLabel("尚未下载视频")
        self.url_file_name.setStyleSheet("font-weight: 500; font-size: 14px;")
        ufi_layout.addWidget(self.url_file_name)
        self.url_file_detail = QLabel("")
        self.url_file_detail.setStyleSheet("color: #86868B; font-size: 12px;")
        ufi_layout.addWidget(self.url_file_detail)
        ufi_layout.addStretch()
        url_layout.addWidget(self.url_file_info)

        # ── 多集链接批量输入区域 ──
        batch_url_frame = QFrame()
        batch_url_frame.setObjectName("batchUrlFrame")
        batch_url_frame.setStyleSheet(
            "QFrame#batchUrlFrame { background-color: #F8F9FE; border: 1px dashed #B0CAFF; "
            "border-radius: 10px; padding: 4px; }"
        )
        bfu_layout = QVBoxLayout(batch_url_frame)
        bfu_layout.setSpacing(8)

        # 标题行
        bfu_header = QHBoxLayout()
        bfu_icon = QLabel("📋")
        bfu_icon.setStyleSheet("font-size: 18px;")
        bfu_header.addWidget(bfu_icon)
        bfu_title = QLabel("多集链接批量分析")
        bfu_title.setStyleSheet("font-size: 14px; font-weight: 600; color: #1D1D1F;")
        bfu_header.addWidget(bfu_title)
        bfu_header.addStretch()
        self.batch_url_count_label = QLabel("")
        self.batch_url_count_label.setStyleSheet("font-size: 11px; color: #86868B;")
        bfu_header.addWidget(self.batch_url_count_label)
        bfu_layout.addLayout(bfu_header)

        # 多行链接输入框
        self.batch_url_input = QTextEdit()
        self.batch_url_input.setPlaceholderText(
            "一行粘贴一个视频链接，例如：\n"
            "https://v.douyin.com/xxxxx\n"
            "https://www.youtube.com/watch?v=xxxxx\n"
            "https://www.bilibili.com/video/BVxxxxx"
        )
        self.batch_url_input.setMaximumHeight(100)
        self.batch_url_input.setStyleSheet(
            "QTextEdit { border: 1px solid #E5E5EA; border-radius: 8px; "
            "padding: 8px 12px; font-size: 12px; background-color: #FFFFFF; "
            "font-family: 'Menlo', 'Consolas', monospace; line-height: 1.5; }"
            "QTextEdit:focus { border-color: #007AFF; }"
        )
        self.batch_url_input.textChanged.connect(self._on_batch_url_text_changed)
        bfu_layout.addWidget(self.batch_url_input)

        # 批量分析按钮
        bfu_btn_row = QHBoxLayout()
        bfu_btn_row.addStretch()
        self.batch_url_analyze_btn = QPushButton("🚀 批量下载分析")
        self.batch_url_analyze_btn.setEnabled(False)
        self.batch_url_analyze_btn.setStyleSheet(
            "QPushButton { background-color: #5856D6; color: white; border: none; "
            "border-radius: 8px; padding: 10px 28px; font-weight: 600; font-size: 14px; }"
            "QPushButton:hover { background-color: #403E99; }"
            "QPushButton:disabled { background-color: #C7C7CC; color: #E5E5EA; }"
        )
        bfu_btn_row.addWidget(self.batch_url_analyze_btn)
        bfu_btn_row.addStretch()
        bfu_layout.addLayout(bfu_btn_row)

        url_layout.addWidget(batch_url_frame)
        self.input_stack.addWidget(url_page)

        content_layout.addWidget(self.input_stack)

        # 通用分析按钮（本地/链接共用）
        self.analyze_btn = QPushButton("▶ 开始分析")
        self.analyze_btn.setObjectName("analyzeBtn")
        self.analyze_btn.setEnabled(False)
        self.analyze_btn.setFixedHeight(42)
        content_layout.addWidget(self.analyze_btn, alignment=Qt.AlignCenter)

        # ── 批量分析工具栏（统一管理所有批量相关按钮）──
        self.batch_toolbar = QFrame()
        self.batch_toolbar.setObjectName("batchToolbar")
        self.batch_toolbar.setVisible(False)
        self.batch_toolbar.setStyleSheet(
            "QFrame#batchToolbar { background-color: #F5F5F7; border-radius: 10px; "
            "border: 1px solid #E5E5EA; padding: 2px; }"
        )
        tb_layout = QHBoxLayout(self.batch_toolbar)
        tb_layout.setContentsMargins(12, 8, 12, 8)
        tb_layout.setSpacing(12)

        # 返回批量视图按钮（单集查看时显示）
        self.back_to_batch_btn = QPushButton("← 返回汇总")
        self.back_to_batch_btn.setVisible(False)
        self.back_to_batch_btn.setStyleSheet(
            "QPushButton { background-color: #FFFFFF; color: #5856D6; border: 1px solid #5856D6; "
            "border-radius: 6px; padding: 6px 14px; font-weight: 500; font-size: 12px; }"
            "QPushButton:hover { background-color: #EDEDFC; }"
        )
        self.back_to_batch_btn.clicked.connect(self._on_back_to_batch)
        tb_layout.addWidget(self.back_to_batch_btn)

        # 集数选择器
        es_label = QLabel("📋 查看集数:")
        es_label.setStyleSheet("font-size: 12px; color: #1D1D1F; font-weight: 500;")
        tb_layout.addWidget(es_label)
        from PySide6.QtWidgets import QComboBox
        self.episode_combo = QComboBox()
        self.episode_combo.setMinimumWidth(200)
        self.episode_combo.setStyleSheet(
            "QComboBox { border: 1px solid #E5E5EA; border-radius: 6px; "
            "padding: 5px 10px; font-size: 12px; background-color: #FFFFFF; }"
            "QComboBox:hover { border-color: #007AFF; }"
        )
        self.episode_combo.currentIndexChanged.connect(self._on_episode_changed)
        tb_layout.addWidget(self.episode_combo, 1)

        tb_layout.addStretch()

        # 全剧总结按钮
        self.summary_btn = QPushButton("📊 生成全剧总结")
        self.summary_btn.setFixedHeight(32)
        self.summary_btn.setStyleSheet(
            "QPushButton { background-color: #5856D6; color: white; border: none; "
            "border-radius: 6px; padding: 6px 16px; font-weight: 600; font-size: 12px; }"
            "QPushButton:hover { background-color: #403E99; }"
            "QPushButton:disabled { background-color: #C7C7CC; color: #E5E5EA; }"
        )
        tb_layout.addWidget(self.summary_btn)

        # 查看已缓存总结按钮
        self.view_summary_btn = QPushButton("📊 查看全剧总结")
        self.view_summary_btn.setVisible(False)
        self.view_summary_btn.setFixedHeight(32)
        self.view_summary_btn.setStyleSheet(
            "QPushButton { background-color: #34C759; color: white; border: none; "
            "border-radius: 6px; padding: 6px 16px; font-weight: 600; font-size: 12px; }"
            "QPushButton:hover { background-color: #2DB84E; }"
        )
        self.view_summary_btn.clicked.connect(self._on_view_cached_summary)
        tb_layout.addWidget(self.view_summary_btn)

        content_layout.addWidget(self.batch_toolbar)

        # 进度
        progress_frame = QFrame()
        pf_layout = QHBoxLayout(progress_frame)
        pf_layout.setContentsMargins(0, 0, 0, 0)
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setFixedHeight(6)
        self.progress_bar.setTextVisible(False)
        pf_layout.addWidget(self.progress_bar)
        self.step_label = QLabel("")
        self.step_label.setObjectName("stepLabel")
        self.step_label.setFixedWidth(280)
        pf_layout.addWidget(self.step_label)
        self.progress_frame = progress_frame
        progress_frame.setVisible(False)
        content_layout.addWidget(progress_frame)

        # ── 展开/收起按钮 ──
        self.expand_btn = QPushButton("▼ 展开分析结果")
        self.expand_btn.setVisible(False)
        self.expand_btn.setStyleSheet(
            "QPushButton { background-color: #F5F5F7; color: #007AFF; border: 1px solid #D1D1D6; "
            "border-radius: 8px; padding: 6px 16px; font-weight: 500; font-size: 13px; }"
            "QPushButton:hover { background-color: #E8F0FE; border-color: #007AFF; }"
        )
        self.expand_btn.setCursor(Qt.PointingHandCursor)
        content_layout.addWidget(self.expand_btn, alignment=Qt.AlignCenter)

        # ── Tab 内容区域 ──
        self.tab_widget = QTabWidget()
        self.tab_widget.setVisible(False)
        self.transcript_tab = QTextEdit()
        self.transcript_tab.setObjectName("transcript")
        self.transcript_tab.setReadOnly(True)
        self.tab_widget.addTab(self.transcript_tab, "转写文本")
        from PySide6.QtWidgets import QTextBrowser
        self.scenes_tab = QTextBrowser()
        self.scenes_tab.setOpenExternalLinks(False)
        self.scenes_tab.setOpenLinks(False)
        self.scenes_tab.anchorClicked.connect(self._on_scene_image_clicked)
        self.tab_widget.addTab(self.scenes_tab, "场景切割")

        # 分析模块：支持双击放大查看
        self.hooks_tab = QTextEdit()
        self.hooks_tab.setReadOnly(True)
        self.hooks_tab.setToolTip("双击可放大查看")
        self.tab_widget.addTab(self.hooks_tab, "钩子分析")
        self.script_tab = QTextEdit()
        self.script_tab.setReadOnly(True)
        self.script_tab.setToolTip("双击可放大查看")
        self.tab_widget.addTab(self.script_tab, "结构化剧本")
        self.characters_tab = QTextEdit()
        self.characters_tab.setReadOnly(True)
        self.characters_tab.setToolTip("双击可放大查看")
        self.tab_widget.addTab(self.characters_tab, "人物图谱")

        # 改写建议 Tab（基于 Excel 爆款短剧分析框架）
        self.rewrite_tab = QTextEdit()
        self.rewrite_tab.setReadOnly(True)
        self.rewrite_tab.setToolTip("双击可放大查看")
        self.tab_widget.addTab(self.rewrite_tab, "✍ 改写建议")

        # 北美改编 Tab（基于 Excel 爆款短剧分析框架 — 5大板块分析）
        self.na_tab = QTextEdit()
        self.na_tab.setReadOnly(True)
        self.na_tab.setToolTip("双击可放大查看")
        self.tab_widget.addTab(self.na_tab, "🌎 北美改编")

        # 双击放大
        self.hooks_tab.mouseDoubleClickEvent = lambda e: self._expand_tab("钩子分析", self.hooks_tab.toMarkdown())
        self.script_tab.mouseDoubleClickEvent = lambda e: self._expand_tab("结构化剧本", self.script_tab.toMarkdown())
        self.characters_tab.mouseDoubleClickEvent = lambda e: self._expand_tab("人物图谱", self.characters_tab.toMarkdown())
        self.rewrite_tab.mouseDoubleClickEvent = lambda e: self._expand_tab("改写建议", self.rewrite_tab.toMarkdown())
        self.na_tab.mouseDoubleClickEvent = lambda e: self._expand_tab("北美改编", self.na_tab.toMarkdown())
        content_layout.addWidget(self.tab_widget, 1)

        # 底部
        bottom = QHBoxLayout()
        bottom.addStretch()
        self.export_md_btn = QPushButton("📄 导出 Markdown")
        self.export_md_btn.setObjectName("exportBtn")
        self.export_md_btn.setEnabled(False)
        bottom.addWidget(self.export_md_btn)
        self.export_docx_btn = QPushButton("📝 导出 Word")
        self.export_docx_btn.setObjectName("exportBtn")
        self.export_docx_btn.setEnabled(False)
        bottom.addWidget(self.export_docx_btn)
        content_layout.addLayout(bottom)

        main_layout.addWidget(content, 1)
        self._update_api_status()

    def _connect_signals(self):
        self.drop_zone.file_dropped.connect(self._on_file_dropped)
        self.browse_btn.clicked.connect(self._on_browse)
        self.analyze_btn.clicked.connect(self._on_analyze)
        self.summary_btn.clicked.connect(self._on_generate_summary)
        self.settings_btn.clicked.connect(self._on_settings)
        self.export_md_btn.clicked.connect(self._on_export_md)
        self.export_docx_btn.clicked.connect(self._on_export_docx)
        self.expand_btn.clicked.connect(self._on_expand_toggle)
        self.input_switch.mode_changed.connect(self._on_mode_changed)
        self.url_zone.download_requested.connect(self._on_download_video)
        self.url_zone.preview_requested.connect(self._on_preview_url)
        self.url_zone.download_btn.clicked.connect(
            lambda: self.url_zone.download_requested.emit(self.url_zone.url_input.text().strip())
        )
        self.url_zone.preview_btn.clicked.connect(
            lambda: self.url_zone.preview_requested.emit(self.url_zone.url_input.text().strip())
        )
        self.batch_url_analyze_btn.clicked.connect(self._on_batch_url_analyze)
        # 线程安全信号 → 槽
        self._progress_signal.connect(self._handle_progress)
        self._analyze_done_signal.connect(self._handle_analyze_done)
        self._download_progress_signal.connect(self._handle_download_progress)
        self._preview_done_signal.connect(self._handle_preview_result)
        self._download_done_signal.connect(self._handle_download_result)
        self._batch_url_progress_signal.connect(self._handle_batch_url_progress)
        self._batch_url_done_signal.connect(self._handle_batch_url_done)

    # ─── 线程安全的 UI 更新槽 ───

    @Slot(int, str)
    def _handle_progress(self, pct, text):
        self.progress_bar.setValue(pct)
        if text:
            self.step_label.setText(text)

    @Slot()
    def _handle_analyze_done(self):
        self.is_analyzing = False
        self.analyze_btn.setEnabled(True)
        self.analyze_btn.setText("▶ 开始分析")

        if self.result and not self.result.error:
            self.progress_bar.setValue(100)
            self.step_label.setText("✅ 分析完成")
            self.step_label.setStyleSheet("color: #34C759; font-size: 12px;")

            # 追加到批量结果（单集分析时也追加）
            if self.view_mode != "batch":
                self.batch_results.append(self.result)

            # 显示批量分析工具栏
            self.batch_toolbar.setVisible(True)
            self._populate_episode_selector()
            self.episode_combo.setCurrentIndex(len(self.batch_results) - 1)

            # 如果有已缓存的总结，显示查看按钮
            if self.cached_summary_md:
                self.view_summary_btn.setVisible(True)

            self._display_results()
            self.export_md_btn.setEnabled(True)
            self.export_docx_btn.setEnabled(True)
            # 显示展开按钮，并自动展开结果区域
            self.expand_btn.setVisible(True)
            self._set_results_expanded(True)
        else:
            self.progress_bar.setValue(0)
            err = self.result.error if self.result else "未知错误"
            self.step_label.setText(f"❌ {err}")
            self.step_label.setStyleSheet("color: #FF3B30; font-size: 12px;")

    def _on_expand_toggle(self):
        """切换结果区域展开/收起"""
        is_visible = self.tab_widget.isVisible()
        self._set_results_expanded(not is_visible)

    def _set_results_expanded(self, expanded: bool):
        """设置结果区域展开状态"""
        self.tab_widget.setVisible(expanded)
        if expanded:
            self.expand_btn.setText("▲ 收起分析结果")
        else:
            self.expand_btn.setText("▼ 展开分析结果")

    @Slot(str)
    def _handle_download_progress(self, msg):
        self.step_label.setText(msg)
        if "获取视频信息" in msg:
            self.progress_bar.setValue(10)
        elif "下载" in msg and "%" in msg:
            try:
                pct = int(msg.split()[-1].replace("%", ""))
                download_pct = int(10 + pct * 0.5)
                self.progress_bar.setValue(download_pct)
            except (ValueError, IndexError):
                pass
        elif "合并" in msg:
            self.progress_bar.setValue(60)

    @Slot(object)
    def _handle_preview_result(self, result):
        self.url_zone.preview_btn.setEnabled(True)
        self.url_zone.preview_btn.setText("预览")
        if result.success:
            self.url_zone.show_preview(result.title, result.platform, result.duration, result.uploader)
        else:
            self.url_zone.preview_frame.setVisible(True)
            self.url_zone.preview_title.setText(f"❌ {result.error}")
            self.url_zone.preview_detail.setText("")

    @Slot(object)
    def _handle_download_result(self, result):
        self.is_downloading = False
        self.url_zone.download_btn.setEnabled(True)
        self.url_zone.download_btn.setText("⬇ 下载")
        if result.success:
            self.video_path = result.video_path
            self.video_source = "url"
            size_mb = os.path.getsize(result.video_path) / (1024 * 1024)
            dur_str = f"{int(result.duration // 60)}:{int(result.duration % 60):02d}" if result.duration else "未知"
            self.url_file_name.setText(f"📹 {result.title}")
            self.url_file_detail.setText(
                f"平台: {result.platform} · 时长: {dur_str} · 大小: {size_mb:.1f}MB"
            )
            self.progress_bar.setValue(65)
            self.step_label.setText("✅ 视频下载完成，可以开始分析")
            self.step_label.setStyleSheet("color: #34C759; font-size: 12px;")
            self.analyze_btn.setEnabled(True)
        else:
            self.url_file_name.setText("❌ 下载失败")
            self.url_file_detail.setText(result.error)
            self.progress_bar.setValue(0)
            self.step_label.setText(f"❌ {result.error}")
            self.step_label.setStyleSheet("color: #FF3B30; font-size: 12px;")

    # ─── 业务逻辑 ───

    def _on_mode_changed(self, mode):
        self.input_stack.setCurrentIndex(0 if mode == "local" else 1)
        if mode == "local":
            self.analyze_btn.setEnabled(bool(self.video_path and self.video_source == "local"))
        else:
            self.analyze_btn.setEnabled(bool(self.video_path and self.video_source == "url"))

    def _update_api_status(self):
        if self.settings.get("openai_api_key"):
            self.api_status_label.setText("✅ AI 分析已就绪")
            self.api_status_label.setStyleSheet("color: #34C759; font-size: 11px;")
        else:
            self.api_status_label.setText("⚠️ 未配置 API Key（仅转写+场景检测）")
            self.api_status_label.setStyleSheet("color: #FF9500; font-size: 11px;")

    def _on_file_dropped(self, path):
        self._set_video_file(path)

    def _on_browse(self):
        path, _ = QFileDialog.getOpenFileName(
            self, "选择视频文件", "",
            "视频文件 (*.mp4 *.mov *.mkv *.avi *.webm);;所有文件 (*)"
        )
        if path:
            self._set_video_file(path)

    def _set_video_file(self, path):
        self.video_path = path
        self.video_source = "local"
        name = os.path.basename(path)
        size = os.path.getsize(path) / (1024 * 1024)
        self.file_name_label.setText(f"📹 {name}")
        self.file_detail_label.setText(f"大小: {size:.1f}MB · 路径: {path}")
        self.analyze_btn.setEnabled(True)

    def _on_preview_url(self, url):
        if not url or self.is_downloading:
            return
        self.url_zone.preview_btn.setEnabled(False)
        self.url_zone.preview_btn.setText("加载中...")

        def worker():
            result = get_video_info_only(url)
            self._preview_done_signal.emit(result)

        threading.Thread(target=worker, daemon=True).start()

    def _on_download_video(self, url):
        if not url or self.is_downloading or self.is_analyzing:
            return
        self.is_downloading = True
        self.url_zone.download_btn.setEnabled(False)
        self.url_zone.download_btn.setText("⏳ 下载中...")
        self.progress_frame.setVisible(True)
        self.progress_bar.setValue(5)
        self.step_label.setText("正在下载视频...")
        self.step_label.setStyleSheet("color: #007AFF; font-size: 12px;")

        cookie_file = self.settings.get("cookie_file", "")

        def worker():
            def progress_cb(msg):
                self._download_progress_signal.emit(msg)
            result = download_video(url, progress_cb=progress_cb, cookie_file=cookie_file)
            self._download_done_signal.emit(result)

        threading.Thread(target=worker, daemon=True).start()

    def _on_settings(self):
        dialog = SettingsDialog(self.settings, self)
        if dialog.exec() == SettingsDialog.Accepted:
            self.settings = dialog.get_settings()
            try:
                save_settings(self.settings)
            except Exception:
                pass  # _accept 里已提示过用户
            self._update_api_status()

    def _on_analyze(self):
        if self.is_analyzing or not self.video_path:
            return
        self.is_analyzing = True
        self.analyze_btn.setEnabled(False)
        self.analyze_btn.setText("⏳ 分析中...")
        self.progress_frame.setVisible(True)
        self.progress_bar.setValue(0)
        self.tab_widget.setVisible(True)

        start_pct = 65 if self.video_source == "url" else 0
        self.progress_bar.setValue(start_pct)

        steps = ["获取视频信息", "提取音频", "Whisper", "转写", "场景检测", "AI分析", "✅", "报告"]

        def progress_cb(msg):
            pct = self.progress_bar.value()
            # 细粒度进度映射
            step_progress = {
                "获取视频信息": 5, "提取音频": 15, "Whisper": 25,
                "语音转写": 40, "场景检测": 55, "AI 剧本分析": 60,
                "角色标注": 65, "钩子分析完成": 72, "结构化剧本完成": 80,
                "人物图谱完成": 88, "改写建议完成": 95,
                "分析完成": 98,
            }
            for key, val in step_progress.items():
                if key in msg:
                    pct = start_pct + int(val * (100 - start_pct) / 100)
                    break
            pct = min(pct, 98)
            self._progress_signal.emit(pct, msg)

        def worker():
            pipeline = VideoToScriptPipeline(
                whisper_model=self.settings.get("whisper_model", "base"),
                scene_threshold=self.settings.get("scene_threshold", 35.0),
                min_scene_duration=self.settings.get("min_scene_duration", 2.0),
                openai_api_key=self.settings.get("openai_api_key", ""),
                openai_model=self.settings.get("openai_model", "gpt-4o-mini"),
                openai_base_url=self.settings.get("openai_base_url", ""),
                language=self.settings.get("language", "") or None,
            )
            self.result = pipeline.run(
                self.video_path,
                progress_cb=progress_cb,
                source=self.video_source,
                platform=getattr(self, '_download_platform', ''),
                video_title=getattr(self, '_download_title', ''),
            )
            self._analyze_done_signal.emit()

        threading.Thread(target=worker, daemon=True).start()

    def _on_scene_image_clicked(self, url):
        """场景截图点击放大"""
        path = url.toLocalFile()
        if not path:
            path = url.toString().replace("file://", "")
        if path and os.path.exists(path):
            dialog = ImageDialog(path, self)
            dialog.exec()

    def _expand_tab(self, title, content):
        """双击放大查看分析内容"""
        if not content or content.strip() == "需要配置 OpenAI API Key":
            return
        dialog = ContentExpandDialog(title, content, is_markdown=True, parent=self)
        dialog.exec()

    def _display_results(self):
        if not self.result:
            return

        # ── 转写文本 Tab ──
        enriched = self.result.enriched_segments
        bgm = self.result.bgm_info or ""

        if enriched:
            # 有富化结果：显示角色标注 + BGM 分离（不再重复显示原始文本）
            html = ""
            if bgm:
                html += f'<div style="background:linear-gradient(135deg,#FFF3CD,#FFE69C);border-radius:10px;padding:12px 16px;margin:4px 0 12px 0;border:1px solid #FFD60A">'
                html += f'<p style="font-size:14px;color:#856404;margin:0">🎵 <b>背景音乐识别</b></p>'
                html += f'<p style="font-size:13px;color:#6C5A0D;margin:4px 0 0 0">{bgm}</p>'
                html += f'</div>'
            elif self.result.bgm_info == "" or "未检测到" in bgm:
                html += f'<div style="background:#F5F5F7;border-radius:10px;padding:12px 16px;margin:4px 0 12px 0;border:1px solid #E5E5EA">'
                html += f'<p style="font-size:13px;color:#86868B;margin:0">🔇 未检测到背景音乐</p>'
                html += f'</div>'

            # 统计摘要卡片
            speakers = set(seg.get("speaker", "未知角色") for seg in enriched if seg.get("type") != "BGM")
            total_segs = len([s for s in enriched if s.get("type") != "BGM"])
            bgm_segs = len([s for s in enriched if s.get("type") == "BGM"])
            html += f'<div style="display:flex;gap:8px;margin-bottom:16px">'
            html += f'<div style="flex:1;background:linear-gradient(135deg,#E8F0FE,#C4D9FD);border-radius:10px;padding:10px 14px;border:1px solid #B0CAFF">'
            html += f'<p style="font-size:11px;color:#007AFF;margin:0">角色数</p>'
            html += f'<p style="font-size:20px;font-weight:700;color:#0055CC;margin:2px 0 0 0">{len(speakers)}</p>'
            html += f'</div>'
            html += f'<div style="flex:1;background:linear-gradient(135deg,#E8FAF0,#B8EDDA);border-radius:10px;padding:10px 14px;border:1px solid #8FD8B8">'
            html += f'<p style="font-size:11px;color:#34C759;margin:0">对白段</p>'
            html += f'<p style="font-size:20px;font-weight:700;color:#1E8E3E;margin:2px 0 0 0">{total_segs}</p>'
            html += f'</div>'
            html += f'<div style="flex:1;background:linear-gradient(135deg,#FFF3E0,#FFE0B2);border-radius:10px;padding:10px 14px;border:1px solid #FFD180">'
            html += f'<p style="font-size:11px;color:#FF9500;margin:0">BGM段</p>'
            html += f'<p style="font-size:20px;font-weight:700;color:#CC7700;margin:2px 0 0 0">{bgm_segs}</p>'
            html += f'</div>'
            html += f'</div>'

            # 角色对话统计表格
            speaker_stats = {}
            for seg in enriched:
                if seg.get("type") == "BGM":
                    continue
                spk = seg.get("speaker", "未知角色")
                if spk not in speaker_stats:
                    speaker_stats[spk] = {"count": 0, "total_time": 0}
                speaker_stats[spk]["count"] += 1
                speaker_stats[spk]["total_time"] += (seg.get("end", 0) - seg.get("start", 0))

            if speaker_stats:
                html += f'<div style="margin:0 0 12px 0;border-radius:10px;overflow:hidden;border:1px solid #E5E5EA;box-shadow:0 1px 4px rgba(0,0,0,0.04)">'
                html += f'<table style="width:100%;border-collapse:collapse;font-size:13px">'
                html += f'<tr style="background:#F5F5F7">'
                html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">角色</th>'
                html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">台词数</th>'
                html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">说话时长</th>'
                html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">活跃度</th>'
                html += f'</tr>'
                max_count = max(s["count"] for s in speaker_stats.values()) if speaker_stats else 1
                for i, (spk, stats) in enumerate(sorted(speaker_stats.items(), key=lambda x: -x[1]["count"])):
                    row_bg = "background:#FFFFFF" if i % 2 == 0 else "background:#FAFAFA"
                    color = self._speaker_color(spk)
                    bar_pct = stats["count"] / max_count * 100
                    html += f'<tr style="{row_bg}">'
                    html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0"><span style="color:{color};font-weight:600">👤 {spk}</span></td>'
                    html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0;color:#1D1D1F;font-weight:500">{stats["count"]} 句</td>'
                    html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0;color:#86868B">{stats["total_time"]:.0f}s</td>'
                    html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0">'
                    html += f'<div style="display:flex;align-items:center;gap:6px">'
                    html += f'<div style="flex:1;background:#F0F0F0;border-radius:4px;height:8px;overflow:hidden"><div style="width:{bar_pct:.0f}%;background:{color};border-radius:4px;height:8px"></div></div>'
                    html += f'<span style="color:{color};font-size:11px;font-weight:500;min-width:36px">{bar_pct:.0f}%</span>'
                    html += f'</div>'
                    html += f'</td>'
                    html += f'</tr>'
                html += f'</table></div>'

            html += '<h3>🎭 带角色标注的台词</h3>'
            current_speaker = None
            for seg in enriched:
                stype = seg.get("type", "DIALOGUE")
                speaker = seg.get("speaker", "未知角色")
                text = seg.get("text", "")
                start = seg.get("start", 0)
                end = seg.get("end", 0)

                if stype == "BGM":
                    html += f'<div style="background:#FFF8E1;border-left:3px solid #FFD60A;border-radius:4px;padding:4px 10px;margin:4px 0">'
                    html += f'<span style="color:#856404;font-style:italic;font-size:12px">🎵 [{start}s-{end}s] {text}</span>'
                    html += f'</div>'
                else:
                    color = self._speaker_color(speaker)
                    if speaker != current_speaker:
                        html += f'<div style="margin-top:12px;padding:6px 12px;background:{color}15;border-left:3px solid {color};border-radius:4px">'
                        html += f'<b style="color:{color};font-size:14px">👤 {speaker}</b>'
                        html += f'</div>'
                        current_speaker = speaker
                    html += f'<div style="margin:4px 0 4px 20px;padding:4px 8px">'
                    html += f'<span style="color:#86868B;font-size:11px;background:#F5F5F7;border-radius:3px;padding:1px 5px">[{start}s-{end}s]</span> '
                    html += f'<span style="font-size:14px">{text}</span>'
                    html += f'</div>'
        else:
            # 无富化结果：只显示带时间戳文本
            html = ""
            # 视频信息摘要
            vi = self.result.video_info
            html += f'<div style="background:linear-gradient(135deg,#F0F5FF,#E0EAFF);border-radius:10px;padding:12px 16px;margin:0 0 12px 0;border:1px solid #B0CAFF">'
            html += f'<p style="font-size:14px;font-weight:600;color:#0055CC;margin:0">📹 转写文本</p>'
            html += f'<p style="font-size:12px;color:#007AFF;margin:4px 0 0 0">共 {len(self.result.transcript_segments)} 段 · 时长 {vi.duration:.0f}秒</p>'
            html += f'</div>'
            for seg in self.result.transcript_segments:
                html += f'<div style="margin:3px 0;padding:4px 8px;background:#FAFAFA;border-radius:4px">'
                html += f'<span style="color:#007AFF;font-weight:600;font-size:12px">[{seg["start"]}s-{seg["end"]}s]</span> '
                html += f'<span style="font-size:14px">{seg["text"]}</span>'
                html += f'</div>'
        self.transcript_tab.setHtml(html)

        # ── 场景切割 Tab ──
        scene_count = len(self.result.scenes)
        shtml = f'<div style="background:linear-gradient(135deg,#E8F0FE,#C4D9FD);border-radius:10px;padding:12px 16px;margin:0 0 16px 0;border:1px solid #B0CAFF">'
        shtml += f'<p style="font-size:15px;font-weight:600;color:#0055CC;margin:0">🎬 场景切割结果</p>'
        shtml += f'<p style="font-size:12px;color:#007AFF;margin:4px 0 0 0">共检测到 <b>{scene_count}</b> 个场景 · 点击截图可放大</p>'
        shtml += f'</div>'

        # 场景总览表格
        if scene_count > 0:
            total_duration = sum(s.duration for s in self.result.scenes)
            shtml += f'<div style="margin:0 0 12px 0;border-radius:10px;overflow:hidden;border:1px solid #B0CAFF;box-shadow:0 1px 4px rgba(0,0,0,0.04)">'
            shtml += f'<table style="width:100%;border-collapse:collapse;font-size:13px">'
            shtml += f'<tr style="background:#007AFF10">'
            shtml += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#007AFF;border-bottom:2px solid #007AFF40">场景</th>'
            shtml += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#007AFF;border-bottom:2px solid #007AFF40">时间范围</th>'
            shtml += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#007AFF;border-bottom:2px solid #007AFF40">时长</th>'
            shtml += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#007AFF;border-bottom:2px solid #007AFF40">占比</th>'
            shtml += f'</tr>'
            for i, s in enumerate(self.result.scenes):
                row_bg = "background:#FFFFFF" if i % 2 == 0 else "background:#F5F8FF"
                pct = f"{s.duration / total_duration * 100:.0f}%" if total_duration > 0 else "-"
                shtml += f'<tr style="{row_bg}">'
                shtml += f'<td style="padding:7px 12px;border-bottom:1px solid #007AFF15"><span style="background:#007AFF;color:white;border-radius:4px;padding:2px 8px;font-weight:600;font-size:12px">{s.index}</span></td>'
                shtml += f'<td style="padding:7px 12px;border-bottom:1px solid #007AFF15;color:#1D1D1F">{s.start}s → {s.end}s</td>'
                shtml += f'<td style="padding:7px 12px;border-bottom:1px solid #007AFF15;color:#1D1D1F;font-weight:500">{s.duration}s</td>'
                shtml += f'<td style="padding:7px 12px;border-bottom:1px solid #007AFF15">'
                # 进度条样式占比
                shtml += f'<div style="display:flex;align-items:center;gap:6px">'
                bar_pct = s.duration / total_duration * 100 if total_duration > 0 else 0
                shtml += f'<div style="flex:1;background:#E8F0FE;border-radius:4px;height:8px;overflow:hidden"><div style="width:{bar_pct:.0f}%;background:#007AFF;border-radius:4px;height:8px"></div></div>'
                shtml += f'<span style="color:#007AFF;font-size:11px;font-weight:500;min-width:36px">{pct}</span>'
                shtml += f'</div>'
                shtml += f'</td>'
                shtml += f'</tr>'
            shtml += f'</table></div>'

        # 场景详细卡片
        for s in self.result.scenes:
            # 场景卡片
            shtml += f'<div style="background:#FFFFFF;border:1px solid #E5E5EA;border-radius:12px;padding:14px;margin:8px 0;box-shadow:0 1px 3px rgba(0,0,0,0.06)">'
            shtml += f'<div style="display:flex;justify-content:space-between;align-items:center">'
            shtml += f'<div>'
            shtml += f'<span style="background:#007AFF;color:white;border-radius:6px;padding:3px 10px;font-weight:600;font-size:13px">场景{s.index}</span>'
            shtml += f' <span style="color:#1D1D1F;font-size:13px;font-weight:500">{s.start}s → {s.end}s</span>'
            shtml += f'</div>'
            shtml += f'<span style="color:#86868B;font-size:12px">时长 {s.duration}s</span>'
            shtml += f'</div>'
            if s.frame_path and os.path.exists(s.frame_path):
                try:
                    img = QImage(s.frame_path)
                    if not img.isNull():
                        from PySide6.QtCore import QBuffer
                        buffer = QBuffer()
                        buffer.open(QBuffer.OpenModeFlag.WriteOnly)
                        img.scaled(360, 202, Qt.KeepAspectRatio).save(buffer, "JPEG", 80)
                        b64 = base64.b64encode(buffer.data().data()).decode()
                        shtml += f'<div style="margin-top:8px">'
                        shtml += f'<a href="file://{s.frame_path}"><img src="data:image/jpeg;base64,{b64}" width="360" style="border-radius:8px;border:1px solid #E5E5EA;cursor:pointer"/></a>'
                        shtml += f'<div style="text-align:center;margin-top:4px"><a href="file://{s.frame_path}" style="color:#007AFF;font-size:11px;text-decoration:none">🔍 点击查看大图</a></div>'
                        shtml += f'</div>'
                except Exception:
                    pass
            shtml += f'</div>'
        self.scenes_tab.setHtml(shtml)

        # ── 钩子分析 Tab ──
        hooks_md = self.result.hooks_analysis or "需要配置 OpenAI API Key"
        if not hooks_md.startswith("⚠️") and not hooks_md.startswith("❌"):
            hooks_html = self._md_to_enriched_html(hooks_md, "🪝 钩子结构分析", "#FF3B30", "#FFE5E5")
            self.hooks_tab.setHtml(hooks_html)
        else:
            self.hooks_tab.setMarkdown(hooks_md)

        # ── 结构化剧本 Tab ──
        script_md = self.result.script_structure or "需要配置 OpenAI API Key"
        if not script_md.startswith("❌"):
            script_html = self._md_to_enriched_html(script_md, "📜 结构化剧本", "#5856D6", "#EDEDFC")
            self.script_tab.setHtml(script_html)
        else:
            self.script_tab.setMarkdown(script_md)

        # ── 人物图谱 Tab ──
        chars_md = self.result.character_map or "需要配置 OpenAI API Key"
        if not chars_md.startswith("❌"):
            chars_html = self._md_to_enriched_html(chars_md, "👥 人物图谱", "#FF9500", "#FFF3E0")
            self.characters_tab.setHtml(chars_html)
        else:
            self.characters_tab.setMarkdown(chars_md)

        # ── 改写建议 Tab ──
        rewrite_md = self.result.rewrite_suggestions or ""
        if rewrite_md and not rewrite_md.startswith("❌") and not rewrite_md.startswith("⚠️"):
            rewrite_html = self._md_to_enriched_html(rewrite_md, "✍ 改写建议", "#5856D6", "#EDEDFC")
            self.rewrite_tab.setHtml(rewrite_html)
        else:
            self.rewrite_tab.setMarkdown(rewrite_md or "暂无改写建议（需要配置 OpenAI API Key）")

        # ── 北美改编 Tab（基于爆款短剧分析框架）──
        na_md = self.result.north_america or ""
        if na_md and not na_md.startswith("❌") and not na_md.startswith("⚠️"):
            na_html = self._md_to_enriched_html(na_md, "🌎 北美改编可行性分析", "#AF52DE", "#F5E8FF")
            self.na_tab.setHtml(na_html)
        else:
            self.na_tab.setMarkdown(na_md or "暂无北美改编分析（需要配置 OpenAI API Key）")

    def _speaker_color(self, speaker):
        """为角色分配固定颜色"""
        colors = ["#007AFF", "#FF3B30", "#34C759", "#FF9500", "#AF52DE", "#5856D6", "#FF2D55", "#5AC8FA"]
        idx = hash(speaker) % len(colors)
        return colors[idx]

    def _md_to_enriched_html(self, md_text: str, title: str, accent_color: str, bg_color: str) -> str:
        """将 Markdown 文本转为视觉更丰富的 HTML 展示，支持表格渲染"""
        import re
        lines = md_text.split("\n")
        html = ""

        # 顶部标题卡片
        html += f'<div style="background:linear-gradient(135deg,{bg_color},{accent_color}15);border-radius:10px;padding:12px 16px;margin:0 0 16px 0;border:1px solid {accent_color}40">'
        html += f'<p style="font-size:16px;font-weight:700;color:{accent_color};margin:0">{title}</p>'
        html += f'<p style="font-size:12px;color:{accent_color}99;margin:4px 0 0 0">双击可放大查看完整内容</p>'
        html += f'</div>'

        in_list = False
        in_table = False
        table_rows = []

        def flush_table():
            nonlocal in_table, table_rows, html
            if not table_rows:
                return
            html += f'<div style="margin:10px 0;border-radius:10px;overflow:hidden;border:1px solid {accent_color}30;box-shadow:0 1px 4px rgba(0,0,0,0.04)">'
            html += f'<table style="width:100%;border-collapse:collapse;font-size:13px">'

            for i, (cells, is_header, is_separator) in enumerate(table_rows):
                if is_separator:
                    continue
                tag = "th" if is_header else "td"
                row_bg = f"background:{accent_color}10" if is_header else ("background:#FFFFFF" if i % 2 == 1 else f"background:{bg_color}60")
                html += f'<tr style="{row_bg}">'
                for cell in cells:
                    cell_text = self._highlight_bold(cell.strip(), accent_color)
                    cell_text = re.sub(r'\[([^\]]+)\]', lambda m: f'<span style="background:{accent_color}20;color:{accent_color};padding:1px 5px;border-radius:3px;font-size:11px">[{m.group(1)}]</span>', cell_text)
                    if is_header:
                        html += f'<{tag} style="padding:8px 12px;text-align:left;font-weight:600;color:{accent_color};border-bottom:2px solid {accent_color}40;white-space:nowrap">{cell_text}</{tag}>'
                    else:
                        html += f'<{tag} style="padding:7px 12px;text-align:left;border-bottom:1px solid {accent_color}15;color:#1D1D1F;max-width:300px;word-break:break-word">{cell_text}</{tag}>'
                html += '</tr>'

            html += '</table></div>'
            table_rows = []
            in_table = False

        for line in lines:
            stripped = line.strip()

            if not stripped:
                if in_list:
                    html += '</div>'
                    in_list = False
                if in_table:
                    flush_table()
                continue

            # 检测 Markdown 表格行
            if '|' in stripped and stripped.count('|') >= 2:
                cells = [c.strip() for c in stripped.split('|')[1:-1]]  # 去掉首尾空元素
                if not cells:
                    cells = [c.strip() for c in stripped.strip('|').split('|')]

                # 判断是否是分隔行（如 |---|---|---|）
                is_separator = all(re.match(r'^[-:]+$', c.strip()) for c in cells if c.strip())

                if not in_table:
                    # 关闭之前的列表
                    if in_list:
                        html += '</div>'
                        in_list = False
                    in_table = True
                    table_rows = []

                is_header = (len(table_rows) == 0)
                table_rows.append((cells, is_header, is_separator))
                continue
            else:
                if in_table:
                    flush_table()

            # H1 标题
            if stripped.startswith("# ") and not stripped.startswith("## "):
                text = stripped[2:]
                html += f'<div style="background:linear-gradient(90deg,{accent_color},{accent_color}CC);border-radius:8px;padding:10px 14px;margin:12px 0 8px 0">'
                html += f'<span style="color:white;font-weight:700;font-size:16px">{text}</span>'
                html += f'</div>'
            # H2 标题
            elif stripped.startswith("## ") and not stripped.startswith("### "):
                text = stripped[3:]
                html += f'<div style="border-left:4px solid {accent_color};padding:8px 12px;margin:10px 0 6px 0;background:{bg_color};border-radius:0 8px 8px 0">'
                html += f'<span style="font-weight:700;font-size:15px;color:{accent_color}">{text}</span>'
                html += f'</div>'
            # H3 标题
            elif stripped.startswith("### "):
                text = stripped[4:]
                html += f'<div style="border-left:3px solid {accent_color}80;padding:6px 10px;margin:8px 0 4px 0;background:{bg_color}60;border-radius:0 6px 6px 0">'
                html += f'<span style="font-weight:600;font-size:14px;color:{accent_color}DD">{text}</span>'
                html += f'</div>'
            # 列表项
            elif stripped.startswith("- ") or stripped.startswith("* "):
                text = stripped[2:]
                text = self._highlight_bold(text, accent_color)
                if not in_list:
                    html += f'<div style="background:#FAFAFA;border-radius:8px;padding:6px 10px;margin:4px 0">'
                    in_list = True
                html += f'<div style="padding:2px 0 2px 14px;border-left:2px solid {accent_color}40">'
                html += f'<span style="color:{accent_color};margin-right:6px">•</span>{text}'
                html += f'</div>'
            # 有序列表
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".、)":
                text = stripped[2:].strip()
                num = stripped[0]
                text = self._highlight_bold(text, accent_color)
                html += f'<div style="background:#FAFAFA;border-radius:6px;padding:4px 10px;margin:2px 0">'
                html += f'<span style="background:{accent_color};color:white;border-radius:50%;width:20px;height:20px;display:inline-block;text-align:center;font-size:11px;font-weight:600;line-height:20px;margin-right:8px">{num}</span>{text}'
                html += f'</div>'
            # 普通段落
            else:
                if in_list:
                    html += '</div>'
                    in_list = False
                text = self._highlight_bold(stripped, accent_color)
                # 检测带括号的情绪标注
                text = re.sub(r'\[([^\]]+)\]', lambda m: f'<span style="background:{accent_color}20;color:{accent_color};padding:1px 6px;border-radius:3px;font-size:12px">[{m.group(1)}]</span>', text)
                html += f'<p style="margin:4px 0;line-height:1.6;font-size:14px">{text}</p>'

        if in_list:
            html += '</div>'
        if in_table:
            flush_table()

        return html

    @staticmethod
    def _highlight_bold(text: str, accent_color: str) -> str:
        """高亮 Markdown 粗体标记"""
        import re
        return re.sub(r'\*\*([^*]+)\*\*', lambda m: f'<b style="color:{accent_color}">{m.group(1)}</b>', text)

    def _get_desktop_path(self):
        """获取桌面路径（兼容中英文系统）"""
        home = os.path.expanduser("~")
        # 优先尝试中文桌面
        desktop_cn = os.path.join(home, "桌面")
        if os.path.isdir(desktop_cn):
            return desktop_cn
        desktop_en = os.path.join(home, "Desktop")
        if os.path.isdir(desktop_en):
            return desktop_en
        return home

    def _on_export_md(self):
        if not self.result or not self.result.full_report:
            return
        desktop = self._get_desktop_path()
        default_path = os.path.join(desktop, "短剧拆解报告.md")
        path, _ = QFileDialog.getSaveFileName(
            self, "导出报告", default_path, "Markdown (*.md)"
        )
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self.result.full_report)
                self.step_label.setText(f"✅ 已导出到桌面: {os.path.basename(path)}")
                self.step_label.setStyleSheet("color: #34C759; font-size: 12px;")
            except Exception as e:
                self.step_label.setText(f"❌ 导出失败: {e}")
                self.step_label.setStyleSheet("color: #FF3B30; font-size: 12px;")

    def _on_export_docx(self):
        if not self.result:
            return
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        desktop = self._get_desktop_path()
        default_path = os.path.join(desktop, "短剧拆解报告.docx")
        path, _ = QFileDialog.getSaveFileName(
            self, "导出 Word", default_path, "Word 文档 (*.docx)"
        )
        if not path:
            return

        doc = Document()
        result = self.result
        vi = result.video_info

        # ─── 标题 ───
        title = doc.add_heading("短剧拆解报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ─── 视频信息 ───
        doc.add_heading("视频信息", level=1)
        info_table = doc.add_table(rows=4, cols=2, style="Light List Accent 1")
        info_data = [
            ("文件", os.path.basename(vi.path)),
            ("时长", f"{vi.duration:.1f}秒"),
            ("分辨率", f"{vi.width}×{vi.height}"),
            ("文件大小", f"{vi.size_mb:.1f}MB"),
        ]
        if vi.source == "url" and vi.platform:
            info_data.append(("来源平台", vi.platform))
        if vi.title:
            info_data.append(("视频标题", vi.title))
        # 动态扩展表格行
        for i, (k, v) in enumerate(info_data):
            if i >= len(info_table.rows):
                info_table.add_row()
            info_table.rows[i].cells[0].text = k
            info_table.rows[i].cells[1].text = v

        # ─── 背景音乐 ───
        if result.bgm_info:
            doc.add_heading("背景音乐", level=1)
            doc.add_paragraph(result.bgm_info)

        # ─── 转写文本 ───
        doc.add_heading("转写文本", level=1)
        enriched = result.enriched_segments
        if enriched:
            current_speaker = None
            for seg in enriched:
                stype = seg.get("type", "DIALOGUE")
                speaker = seg.get("speaker", "未知角色")
                text = seg.get("text", "")
                start = seg.get("start", 0)
                end = seg.get("end", 0)
                time_tag = f"[{start}s-{end}s]"

                if stype == "BGM":
                    p = doc.add_paragraph()
                    run = p.add_run(f"🎵 {time_tag} {text}")
                    run.font.color.rgb = RGBColor(0xAE, 0xAE, 0xB2)
                    run.italic = True
                else:
                    if speaker != current_speaker:
                        p = doc.add_paragraph()
                        run = p.add_run(f"👤 {speaker}")
                        run.bold = True
                        color_map = {
                            "007AFF": RGBColor(0x00, 0x7A, 0xFF),
                            "FF3B30": RGBColor(0xFF, 0x3B, 0x30),
                            "34C759": RGBColor(0x34, 0xC7, 0x59),
                            "FF9500": RGBColor(0xFF, 0x95, 0x00),
                            "AF52DE": RGBColor(0xAF, 0x52, 0xDE),
                        }
                        c = self._speaker_color(speaker).lstrip("#")
                        run.font.color.rgb = color_map.get(c, RGBColor(0x00, 0x7A, 0xFF))
                        current_speaker = speaker
                    p = doc.add_paragraph(f"    {time_tag} {text}")
        else:
            for seg in result.transcript_segments:
                doc.add_paragraph(f"[{seg['start']}s-{seg['end']}s] {seg['text']}")

        # ─── 场景切割 ───
        doc.add_heading("场景切割", level=1)
        for s in result.scenes:
            p = doc.add_paragraph()
            run = p.add_run(f"场景{s.index}")
            run.bold = True
            p.add_run(f" · {s.start}s → {s.end}s (时长{s.duration}s)")
            if s.frame_path and os.path.exists(s.frame_path):
                try:
                    doc.add_picture(s.frame_path, width=Inches(4.5))
                except Exception:
                    pass

        # ─── 钩子分析 ───
        if result.hooks_analysis and not result.hooks_analysis.startswith("⚠️") and not result.hooks_analysis.startswith("❌"):
            doc.add_heading("钩子结构分析", level=1)
            self._add_markdown_to_doc(doc, result.hooks_analysis)

        # ─── 结构化剧本 ───
        if result.script_structure and not result.script_structure.startswith("❌"):
            doc.add_heading("结构化剧本", level=1)
            self._add_markdown_to_doc(doc, result.script_structure)

        # ─── 人物图谱 ───
        if result.character_map and not result.character_map.startswith("❌"):
            doc.add_heading("人物图谱", level=1)
            self._add_markdown_to_doc(doc, result.character_map)

        # ─── 改写建议 ───
        if result.rewrite_suggestions and not result.rewrite_suggestions.startswith("❌"):
            doc.add_heading("改写建议", level=1)
            self._add_markdown_to_doc(doc, result.rewrite_suggestions)

        # ─── 北美改编分析 ───
        if result.north_america and not result.north_america.startswith("❌"):
            doc.add_heading("北美改编可行性分析", level=1)
            self._add_markdown_to_doc(doc, result.north_america)

        try:
            doc.save(path)
            self.step_label.setText(f"✅ 已导出 Word 到桌面: {os.path.basename(path)}")
            self.step_label.setStyleSheet("color: #34C759; font-size: 12px;")
        except Exception as e:
            self.step_label.setText(f"❌ Word导出失败: {e}")
            self.step_label.setStyleSheet("color: #FF3B30; font-size: 12px;")

    @staticmethod
    def _add_markdown_to_doc(doc, md_text):
        """简易 Markdown → Word 段落转换"""
        for line in md_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".、)":
                doc.add_paragraph(stripped[2:].strip(), style="List Number")
            else:
                p = doc.add_paragraph()
                # 处理粗体 **text**
                parts = stripped.split("**")
                for i, part in enumerate(parts):
                    if not part:
                        continue
                    run = p.add_run(part)
                    if i % 2 == 1:  # 奇数部分是粗体内容
                        run.bold = True

    # ─── 多集链接批量下载分析 ───

    def _on_batch_url_text_changed(self):
        """多集链接输入框文本变化，更新链接计数"""
        urls = self._parse_batch_urls()
        count = len(urls)
        if count > 0:
            self.batch_url_count_label.setText(f"已输入 {count} 个有效链接")
            self.batch_url_analyze_btn.setEnabled(True)
        else:
            self.batch_url_count_label.setText("")
            self.batch_url_analyze_btn.setEnabled(False)

    def _parse_batch_urls(self):
        """解析多集链接输入框中的 URL 列表"""
        text = self.batch_url_input.toPlainText().strip()
        if not text:
            return []
        urls = []
        for line in text.split("\n"):
            line = line.strip()
            if line and line.startswith("http"):
                urls.append(line)
        return urls

    def _on_batch_url_analyze(self):
        """多集链接批量下载分析：逐个下载 + 分析"""
        urls = self._parse_batch_urls()
        if not urls or self.is_batch_url_analyzing or self.is_analyzing:
            return

        self.is_batch_url_analyzing = True
        self.batch_url_analyze_btn.setEnabled(False)
        self.batch_url_analyze_btn.setText("⏳ 批量分析中...")
        self.batch_url_input.setReadOnly(True)

        # 清空之前的批量结果
        self.batch_results = []
        self.view_mode = "batch"

        self.progress_frame.setVisible(True)
        self.progress_bar.setValue(0)
        self.step_label.setText(f"多集批量分析：0/{len(urls)} 准备中...")
        self.step_label.setStyleSheet("color: #5856D6; font-size: 12px;")

        cookie_file = self.settings.get("cookie_file", "")

        def worker():
            total = len(urls)
            for i, url in enumerate(urls):
                # 报告当前进度
                self._batch_url_progress_signal.emit(i + 1, total, f"正在下载第 {i+1}/{total} 集...")

                # 第一步：下载
                try:
                    dl_result = download_video(url, progress_cb=lambda msg: None, cookie_file=cookie_file)
                    if not dl_result.success:
                        # 下载失败，跳过这集，记录错误
                        err_result = AnalysisResult(
                            video_info=None,
                            transcript_segments=[],
                            enriched_segments=None,
                            scenes=[],
                            hooks_analysis="",
                            script="",
                            character_map="",
                            rewrite_suggestions="",
                            north_america="",
                            bgm_info="",
                            error=f"第{i+1}集下载失败: {dl_result.error}"
                        )
                        self.batch_results.append(err_result)
                        continue

                    video_path = dl_result.video_path
                    platform = dl_result.platform
                    title = dl_result.title
                except Exception as e:
                    err_result = AnalysisResult(
                        video_info=None,
                        transcript_segments=[],
                        enriched_segments=None,
                        scenes=[],
                        hooks_analysis="",
                        script="",
                        character_map="",
                        rewrite_suggestions="",
                        north_america="",
                        bgm_info="",
                        error=f"第{i+1}集下载异常: {e}"
                    )
                    self.batch_results.append(err_result)
                    continue

                # 第二步：分析
                self._batch_url_progress_signal.emit(i + 1, total, f"第 {i+1}/{total} 集分析中...")

                try:
                    pipeline = VideoToScriptPipeline(
                        whisper_model=self.settings.get("whisper_model", "base"),
                        scene_threshold=self.settings.get("scene_threshold", 35.0),
                        min_scene_duration=self.settings.get("min_scene_duration", 2.0),
                        openai_api_key=self.settings.get("openai_api_key", ""),
                        openai_model=self.settings.get("openai_model", "gpt-4o-mini"),
                        openai_base_url=self.settings.get("openai_base_url", ""),
                        language=self.settings.get("language", "") or None,
                    )
                    result = pipeline.run(
                        video_path,
                        progress_cb=lambda msg: None,
                        source="url",
                        platform=platform,
                        video_title=title,
                    )
                    self.batch_results.append(result)
                except Exception as e:
                    err_result = AnalysisResult(
                        video_info=None,
                        transcript_segments=[],
                        enriched_segments=None,
                        scenes=[],
                        hooks_analysis="",
                        script="",
                        character_map="",
                        rewrite_suggestions="",
                        north_america="",
                        bgm_info="",
                        error=f"第{i+1}集分析异常: {e}"
                    )
                    self.batch_results.append(err_result)

            # 全部完成
            self._batch_url_done_signal.emit()

        threading.Thread(target=worker, daemon=True).start()

    @Slot(int, int, str)
    def _handle_batch_url_progress(self, current, total, msg):
        """多集批量分析进度更新"""
        pct = int(current / total * 100) if total > 0 else 0
        self.progress_bar.setValue(pct)
        self.step_label.setText(msg)
        self.step_label.setStyleSheet("color: #5856D6; font-size: 12px;")

    @Slot()
    def _handle_batch_url_done(self):
        """多集批量分析完成"""
        self.is_batch_url_analyzing = False
        self.batch_url_analyze_btn.setEnabled(True)
        self.batch_url_analyze_btn.setText("🚀 批量下载分析")
        self.batch_url_input.setReadOnly(False)

        success_count = sum(1 for r in self.batch_results if not r.error)
        fail_count = sum(1 for r in self.batch_results if r.error)

        self.progress_bar.setValue(100)
        if fail_count == 0:
            self.step_label.setText(f"✅ 批量分析完成！共 {success_count} 集全部成功")
            self.step_label.setStyleSheet("color: #34C759; font-size: 12px;")
        else:
            self.step_label.setText(f"⚠️ 批量分析完成：{success_count} 集成功，{fail_count} 集失败")
            self.step_label.setStyleSheet("color: #FF9500; font-size: 12px;")

        # 显示集数选择器和总结按钮
        if len(self.batch_results) > 0:
            self.tab_widget.setVisible(True)
            self.expand_btn.setVisible(True)
            self._set_results_expanded(True)

            # 如果有成功的，设置当前结果为第一个成功的
            for r in self.batch_results:
                if not r.error:
                    self.result = r
                    break

            # 显示批量分析工具栏
            self.batch_toolbar.setVisible(True)
            self._populate_episode_selector()
            if len(self.batch_results) > 0:
                self.episode_combo.setCurrentIndex(0)

            self.export_md_btn.setEnabled(True)
            self.export_docx_btn.setEnabled(True)

            # 显示批量汇总视图
            self._display_batch_results()

    # ─── 批量分析与多集管理 ───

    def _on_back_to_batch(self):
        """返回批量结果视图"""
        if not self.batch_results:
            return
        self.view_mode = "batch"
        self.back_to_batch_btn.setVisible(False)
        self._display_batch_results()

    def _on_episode_changed(self, index):
        """集数选择器切换"""
        if index < 0 or index >= len(self.batch_results):
            return
        self.current_episode_idx = index
        self.view_mode = "single"
        self.result = self.batch_results[index]
        self._display_results()
        # 单集视图时显示返回按钮
        self.back_to_batch_btn.setVisible(len(self.batch_results) > 1)

    def _populate_episode_selector(self):
        """填充集数选择器"""
        self.episode_combo.blockSignals(True)
        self.episode_combo.clear()
        for i, r in enumerate(self.batch_results):
            vi = r.video_info
            title = vi.title if vi and vi.title else os.path.basename(vi.path) if vi else f"第{i+1}集"
            if len(title) > 30:
                title = title[:30] + "..."
            self.episode_combo.addItem(f"第{i+1}集: {title}")
        self.episode_combo.blockSignals(False)

    def _display_batch_results(self):
        """显示批量分析结果汇总"""
        if not self.batch_results:
            return
        self.tab_widget.setVisible(True)
        self.expand_btn.setVisible(True)
        self._set_results_expanded(True)

        # 在转写文本 Tab 中显示批量汇总
        html = ""
        html += f'<div style="background:linear-gradient(135deg,#E8F0FE,#C4D9FD);border-radius:10px;padding:14px 16px;margin:0 0 16px 0;border:1px solid #B0CAFF">'
        html += f'<p style="font-size:16px;font-weight:700;color:#0055CC;margin:0">📦 批量分析结果</p>'
        html += f'<p style="font-size:13px;color:#007AFF;margin:4px 0 0 0">共分析 <b>{len(self.batch_results)}</b> 个视频</p>'
        html += f'</div>'

        # 汇总表格
        html += f'<div style="margin:0 0 12px 0;border-radius:10px;overflow:hidden;border:1px solid #E5E5EA;box-shadow:0 1px 4px rgba(0,0,0,0.04)">'
        html += f'<table style="width:100%;border-collapse:collapse;font-size:13px">'
        html += f'<tr style="background:#F5F5F7">'
        html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">集数</th>'
        html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">时长</th>'
        html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">对白段</th>'
        html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">场景数</th>'
        html += f'<th style="padding:8px 12px;text-align:left;font-weight:600;color:#1D1D1F;border-bottom:2px solid #E5E5EA">状态</th>'
        html += f'</tr>'
        for i, r in enumerate(self.batch_results):
            vi = r.video_info
            row_bg = "background:#FFFFFF" if i % 2 == 0 else "background:#FAFAFA"
            dur = f"{vi.duration:.0f}s" if vi else "-"
            segs = len(r.enriched_segments) if r.enriched_segments else len(r.transcript_segments)
            scenes = len(r.scenes)
            status = "✅" if not r.error else "❌"
            html += f'<tr style="{row_bg}">'
            html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0;font-weight:500">第{i+1}集</td>'
            html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0;color:#86868B">{dur}</td>'
            html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0">{segs}</td>'
            html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0">{scenes}</td>'
            html += f'<td style="padding:7px 12px;border-bottom:1px solid #F0F0F0">{status}</td>'
            html += f'</tr>'
        html += f'</table></div>'
        html += f'<p style="color:#86868B;font-size:12px;margin:8px 0">💡 点击上方「查看集数」下拉框切换到具体集数查看详细分析</p>'
        self.transcript_tab.setHtml(html)

        # 其他 Tab 显示提示
        placeholder = '<p style="color:#86868B;font-size:14px;text-align:center;margin-top:40px">请选择具体集数查看详细分析</p>'
        self.scenes_tab.setHtml(placeholder)
        self.hooks_tab.setHtml(placeholder)
        self.script_tab.setHtml(placeholder)
        self.characters_tab.setHtml(placeholder)
        self.rewrite_tab.setHtml(placeholder)
        self.na_tab.setHtml(placeholder)

    # ─── 全剧总结 ───

    def _on_generate_summary(self):
        """生成全剧总结（如果已有缓存则直接显示）"""
        if not self.batch_results or not self.settings.get("openai_api_key"):
            self.step_label.setText("❌ 需要批量分析结果和 API Key 才能生成全剧总结")
            self.step_label.setStyleSheet("color: #FF3B30; font-size: 12px;")
            return

        # 如果已有缓存，直接显示（不需要重新调用 LLM）
        if self.cached_summary_md:
            self._on_view_cached_summary()
            return

        self.summary_btn.setEnabled(False)
        self.summary_btn.setText("⏳ 生成总结中...")
        self.view_mode = "summary"

        def worker():
            try:
                from openai import OpenAI
                kwargs = {"api_key": self.settings.get("openai_api_key", ""), "timeout": 180.0}
                base_url = self.settings.get("openai_base_url", "")
                if base_url:
                    kwargs["base_url"] = base_url
                client = OpenAI(**kwargs)

                # 汇总所有集的详细数据（更充分地传入原始内容）
                episode_details = []
                for i, r in enumerate(self.batch_results):
                    if r.error:
                        episode_details.append(f"第{i+1}集: 分析失败({r.error})")
                        continue
                    vi = r.video_info
                    dur = f"{vi.duration:.0f}s" if vi else "未知"
                    ep_title = vi.title if vi and vi.title else f"第{i+1}集"

                    # 转写文本（截取前1500字，保留关键对白）
                    transcript_parts = []
                    if r.enriched_segments:
                        for seg in r.enriched_segments[:30]:
                            speaker = getattr(seg, 'speaker', '') or ''
                            text = getattr(seg, 'text', '') or ''
                            if speaker and text:
                                transcript_parts.append(f"{speaker}：{text}")
                            elif text:
                                transcript_parts.append(text)
                    transcript_text = "\n".join(transcript_parts)[:1500]

                    # 钩子分析
                    hooks = r.hooks_analysis[:800] if r.hooks_analysis else "无"

                    # 人物图谱
                    chars = r.character_map[:800] if r.character_map else "无"

                    # 结构化剧本（含场景描述）
                    script = r.script_structure[:1200] if r.script_structure else "无"

                    # 改写建议
                    rewrite = r.rewrite_suggestions[:600] if r.rewrite_suggestions else "无"

                    detail = (
                        f"## 第{i+1}集《{ep_title}》（时长{dur}）\n"
                        f"### 转写对白\n{transcript_text}\n\n"
                        f"### 钩子分析\n{hooks}\n\n"
                        f"### 人物图谱\n{chars}\n\n"
                        f"### 结构化剧本\n{script}\n\n"
                        f"### 改写建议\n{rewrite}"
                    )
                    episode_details.append(detail)

                all_episodes = "\n\n---\n\n".join(episode_details)

                prompt = f"""你是专业的短剧全剧分析专家。以下是一部短剧各集的完整分析数据，请生成一份「全剧总结报告」。

{all_episodes}

---

请按以下结构输出全剧总结（必须严格按此格式）：

# 全剧总结报告

---

## 第一部分：逐集详细故事描述

对每一集，严格按以下格式输出：

### 第X集《集名》
**故事概况**：这一集讲了什么（3-5句话概括核心剧情）
**关键情节**：按时间顺序列出本集发生的所有关键事件（至少5个情节点）
**人物状态**：本集出现的主要人物、他们的目标/情绪/处境
**钩子与转折**：本集开头用了什么钩子、结尾在哪里断住
**情绪曲线**：本集的情绪走向

---

## 第二部分：全剧综合分析

### 一、全剧爆款公式
用一句话总结整部剧的爆款逻辑

### 二、人物弧光总结
- 主角成长轨迹
- 反派结局
- 核心人物关系变化

### 三、钩子体系评估
- 各集钩子类型分布（用表格）
- 最强钩子与最弱钩子
- 付费墙位置建议（第3/6/10集结尾）

### 四、情绪曲线总览
- 全剧情绪起伏（用表格展示各集情绪关键词）
- 最强的3个情绪高潮点
- 情绪断裂点（如有）

### 五、改编建议
- 最有改编价值的3集
- 整体改写方向推荐
- 北美市场适配要点

用 Markdown 表格和结构化文字输出，确保逐集描述丰富详尽。"""

                resp = client.chat.completions.create(
                    model=self.settings.get("openai_model", "gpt-4o-mini"),
                    messages=[
                        {"role": "system", "content": "你是短剧全剧分析专家，擅长逐集故事解读、多集连贯性分析和改编策略。你的逐集描述必须详尽、具体、基于实际对白内容，不能泛泛而谈。"},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7, max_tokens=8000
                )
                summary_md = resp.choices[0].message.content

                # 缓存到本地
                cache_dir = os.path.expanduser("~/.video_to_script_cache")
                os.makedirs(cache_dir, exist_ok=True)
                series_name = ""
                if self.batch_results and self.batch_results[0].video_info:
                    series_name = self.batch_results[0].video_info.title or os.path.basename(self.batch_results[0].video_info.path)
                series_dir = os.path.join(cache_dir, series_name[:50].replace("/", "_"))
                os.makedirs(series_dir, exist_ok=True)
                with open(os.path.join(series_dir, "_summary.json"), "w", encoding="utf-8") as f:
                    json.dump({"summary": summary_md}, f, ensure_ascii=False, indent=2)

                # 在主线程更新 UI
                QMetaObject.invokeMethod(self, "_show_summary_result", Qt.QueuedConnection, Q_ARG(str, summary_md))

            except Exception as e:
                QMetaObject.invokeMethod(self, "_show_summary_error", Qt.QueuedConnection, Q_ARG(str, str(e)))

        threading.Thread(target=worker, daemon=True).start()

    @Slot(str)
    def _show_summary_result(self, summary_md: str):
        """显示全剧总结结果"""
        # 缓存总结到界面（不丢失）
        self.cached_summary_md = summary_md

        self.summary_btn.setEnabled(True)
        self.summary_btn.setText("📊 重新生成总结")
        # 显示查看缓存总结按钮
        self.view_summary_btn.setVisible(True)

        self.tab_widget.setVisible(True)
        self._set_results_expanded(True)

        # 在转写文本 Tab 显示总结
        summary_html = self._md_to_enriched_html(summary_md, "📊 全剧总结", "#5856D6", "#EDEDFC")
        self.transcript_tab.setHtml(summary_html)

        # 更新其他 Tab 显示提示
        placeholder = '<p style="color:#86868B;font-size:14px;text-align:center;margin-top:40px">📊 当前为全剧总结视图，选择具体集数可查看详细分析</p>'
        self.scenes_tab.setHtml(placeholder)
        self.hooks_tab.setHtml(placeholder)
        self.script_tab.setHtml(placeholder)
        self.characters_tab.setHtml(placeholder)
        self.rewrite_tab.setHtml(placeholder)
        self.na_tab.setHtml(placeholder)

        self.step_label.setText("✅ 全剧总结生成完成")
        self.step_label.setStyleSheet("color: #34C759; font-size: 12px;")

    def _on_view_cached_summary(self):
        """查看已缓存的全剧总结"""
        if not self.cached_summary_md:
            return
        self.view_mode = "summary"
        self.tab_widget.setVisible(True)
        self._set_results_expanded(True)

        summary_html = self._md_to_enriched_html(self.cached_summary_md, "📊 全剧总结", "#5856D6", "#EDEDFC")
        self.transcript_tab.setHtml(summary_html)

        placeholder = '<p style="color:#86868B;font-size:14px;text-align:center;margin-top:40px">📊 当前为全剧总结视图，选择具体集数可查看详细分析</p>'
        self.scenes_tab.setHtml(placeholder)
        self.hooks_tab.setHtml(placeholder)
        self.script_tab.setHtml(placeholder)
        self.characters_tab.setHtml(placeholder)
        self.rewrite_tab.setHtml(placeholder)
        self.na_tab.setHtml(placeholder)

    @Slot(str)
    def _show_summary_error(self, error: str):
        """显示全剧总结错误"""
        self.summary_btn.setEnabled(True)
        self.summary_btn.setText("📊 生成全剧总结")
        self.step_label.setText(f"❌ 全剧总结生成失败: {error}")
        self.step_label.setStyleSheet("color: #FF3B30; font-size: 12px;")

    # 主窗口级拖拽兜底
    def _resolve_path(self, mime):
        if mime.hasUrls():
            url = mime.urls()[0]
            path = url.toLocalFile()
            if not path:
                url_str = url.toString()
                if url_str.startswith("file://"):
                    from urllib.parse import unquote, urlparse
                    path = unquote(urlparse(url_str).path)
            return path
        if mime.hasText():
            return mime.text().strip()
        return ""

    def _is_video(self, path):
        if not path:
            return False
        ext = os.path.splitext(path)[1].lower()
        return ext in (".mp4", ".mov", ".mkv", ".avi", ".webm", ".flv", ".wmv", ".ts")

    def dragEnterEvent(self, event):
        if self._is_video(self._resolve_path(event.mimeData())):
            event.acceptProposedAction()

    def dragMoveEvent(self, event):
        event.acceptProposedAction()

    def dropEvent(self, event):
        path = self._resolve_path(event.mimeData())
        if path:
            self._set_video_file(path)
